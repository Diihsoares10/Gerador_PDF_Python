from docx import Document
from docx2pdf import convert
import os
from database import Database
import datetime
import sys


def substituir_texto_documento(documento, referencias):
    for paragrafo in documento.paragraphs:
        for codigo, valor in referencias.items():
            paragrafo.text = paragrafo.text.replace(codigo, str(valor))
            paragrafo.add_run('').bold = False
            
        
def substituir_texto_tabela(tabela, referencias):
    for linha in tabela.rows:
        for celula in linha.cells:
            for paragrafo in celula.paragraphs:
                for codigo, valor in referencias.items():
                    if paragrafo.text != '':
                        paragrafo.text = paragrafo.text.replace(codigo, str(valor))

def main():
    id = 1
        
    if isinstance(id, int):
        print("Inicio do Processamento...")
            
    db = Database()

    query = f'''SELECT 
                    t1.exemplo1,
                    t2.exemplo2,
                    t3.exemplo3
                       
                FROM 
                    tabela-01 t1 
                    LEFT JOIN tabela-02.coluna t2 ON (t1.coluna = t2.id)
                    LEFT JOIN tabela-03.coluna2 t3 ON (t2.coluna3 = t3.coluna3) 
                    WHERE t1.id = {id}'''

    base = db.read(query)

    # Caminho completo para o documento .docx
    exemplo_pdf = "ORIENTAÇÕES PARA AUDIÊNCIA.docx"

    # Manipular o documento .docx
    documento = Document(exemplo_pdf)

    # Dicionário com todas as referências
    referencias = {
        '<<exemplo_doc>>': base.exemplo1.values[0],
        
    }

    # Lista para armazenar temporariamente os valores verdadeiros
    valores_verdadeiros = []

    # Verifica se base.exemplo2 é True e adiciona à lista
    if base.exemplo2.values[0]:
        valores_verdadeiros.append("Esta ok")
   

    # Verifica se a lista de valores verdadeiros não está vazia
    if valores_verdadeiros:
        # Se houver valores verdadeiros, une-os em uma única string e atribui à chave
        referencias['<<exemplo pra ir adicionando>>'] = ", ".join(valores_verdadeiros)
    else:
        # Se não houver valores verdadeiros, atribui outro valor à chave
        referencias['<<exemplo pra ir adicionando>>'] = ""

    #Verifica se exemplo3 é True e atribui "Sim" ou "Não" conforme a condição
    referencias['<<exemplo3>>'] = "Sim" if base.exemplo3.values[0] else "Não"



    # Substituir texto dentro das tabelas
    for tabela in documento.tables:
        substituir_texto_tabela(tabela, referencias)

    # Substituir texto fora das tabelas
    substituir_texto_documento(documento, referencias)

    # Obter a data e hora atual para uso no nome do arquivo
    now = datetime.datetime.now()

    # Formatar os valores para uso no nome do arquivo
    nome_do_arquivo = f'DOCUMENTO_EXEMPLO.docx'

    # Substituir espaços em branco por sublinhados no nome do arquivo
    nome_do_arquivo = nome_do_arquivo.replace(" ", "_")

    # Salvar o documento com o novo nome de arquivo
    documento.save(nome_do_arquivo)

    # Converter o arquivo modificado para PDF
    pdf_nome = nome_do_arquivo.replace(".docx", ".pdf")
    convert(nome_do_arquivo, pdf_nome)

    # Remover o arquivo temporário .docx
    os.remove(nome_do_arquivo)
    print("Processamento Finalizado.")

